import os
import sys
import glob
import warnings
import base64
from typing import List
from dotenv import load_dotenv

from langchain_community.document_loaders import PyPDFLoader, CSVLoader, UnstructuredWordDocumentLoader
from langchain_core.prompts import PromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma, FAISS
from langchain_core.documents import Document

warnings.filterwarnings("ignore")

# Ensure src path is included (optional, only if src directory exists)
src_paths = ['./src', '../src/']
for path in src_paths:
  if os.path.exists(path):
    sys.path.insert(1, path)

load_dotenv()

# Load API key from environment (don't raise at import time - check in functions)
GEMINI_API_KEY = os.environ.get("GOOGLE_API_KEY")


def load_model():
  """
  Load LLM and embeddings
  """
  if not GEMINI_API_KEY:
    raise ValueError("GOOGLE_API_KEY is not set in environment variables. Please set it in Render environment variables.")
  
  model = ChatGoogleGenerativeAI(
    model=os.environ.get("GEMINI_CHAT_MODEL", "models/gemini-2.5-flash"),
    google_api_key=GEMINI_API_KEY,
    temperature=0.4,
    convert_system_message_to_human=True
  )
  embeddings = GoogleGenerativeAIEmbeddings(
    model=os.environ.get("GEMINI_EMBED_MODEL", "models/text-embedding-004"),
    google_api_key=GEMINI_API_KEY
  )
  return model, embeddings


def extract_text_from_image(image_path: str):
  """
  Extract text from image using Google Gemini Vision API
  """
  try:
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain_core.messages import HumanMessage
    from PIL import Image
    
    # Read and prepare image
    image = Image.open(image_path)
    
    # Use Gemini Vision to extract text
    llm = ChatGoogleGenerativeAI(
      model=os.environ.get("GEMINI_CHAT_MODEL", "models/gemini-2.5-flash"),
      google_api_key=GEMINI_API_KEY,
      temperature=0.1
    )
    
    # Create message with image using langchain's format
    # Gemini supports images directly in HumanMessage content
    message = HumanMessage(
      content=[
        {
          "type": "text",
          "text": "Extract all text from this image. Return only the text content, no explanations."
        },
        {
          "type": "image_url",
          "image_url": image_path
        }
      ]
    )
    
    # Alternative: use the image object directly if supported
    try:
      response = llm.invoke([message])
    except:
      # Fallback: encode as base64
      with open(image_path, 'rb') as img_file:
        image_bytes = img_file.read()
      image_b64 = base64.b64encode(image_bytes).decode('utf-8')
      
      ext = os.path.splitext(image_path)[1].lower()
      mime_types = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
        '.gif': 'image/gif', '.bmp': 'image/bmp', '.webp': 'image/webp'
      }
      mime_type = mime_types.get(ext, 'image/jpeg')
      
      message = HumanMessage(
        content=[
          "Extract all text from this image. Return only the text content, no explanations.",
          {
            "type": "image_url",
            "image_url": f"data:{mime_type};base64,{image_b64}"
          }
        ]
      )
      response = llm.invoke([message])
    
    text_content = response.content if hasattr(response, 'content') else str(response)
    
    # Create a Document from the extracted text
    return [Document(page_content=text_content, metadata={"source": image_path, "type": "image"})]
  except Exception as e:
    print(f"[ERROR] Failed to extract text from image: {e}")
    import traceback
    print(f"[ERROR] Traceback: {traceback.format_exc()}")
    raise ValueError(f"Failed to extract text from image: {str(e)}")


def load_documents(source_dir: str):
  """
  Load documents from multiple sources: PDF, DOC, DOCX, CSV, and images (JPG, PNG, etc.)
  """
  documents = []

  # Supported file extensions
  image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
  doc_extensions = ['.doc', '.docx']
  pdf_extensions = ['.pdf']
  csv_extensions = ['.csv']

  if os.path.isfile(source_dir):
    ext = os.path.splitext(source_dir)[1].lower()
    
    if ext in pdf_extensions:
      documents.extend(PyPDFLoader(source_dir).load())
    elif ext in csv_extensions:
      documents.extend(CSVLoader(source_dir).load())
    elif ext in doc_extensions:
      try:
        # Try UnstructuredWordDocumentLoader first
        documents.extend(UnstructuredWordDocumentLoader(source_dir).load())
      except Exception as e:
        print(f"[WARNING] UnstructuredWordDocumentLoader failed: {e}")
        # Fallback: try python-docx if available
        try:
          import docx
          doc = docx.Document(source_dir)
          text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
          documents.append(Document(page_content=text, metadata={"source": source_dir, "type": "docx"}))
        except ImportError:
          raise ValueError("python-docx is required for DOCX files. Install it with: pip install python-docx")
        except Exception as e2:
          raise ValueError(f"Failed to load DOCX file: {str(e2)}")
    elif ext in image_extensions:
      documents.extend(extract_text_from_image(source_dir))
    else:
      raise ValueError(f"Unsupported file type: {ext}. Supported types: PDF, DOC, DOCX, CSV, JPG, PNG, GIF, BMP, WEBP")
  else:
    # Directory mode - load all supported files
    for pattern in ["*.pdf", "*.doc", "*.docx", "*.csv", "*.jpg", "*.jpeg", "*.png", "*.gif", "*.bmp", "*.webp"]:
      for file_path in glob.glob(os.path.join(source_dir, pattern)):
        ext = os.path.splitext(file_path)[1].lower()
        if ext in pdf_extensions:
          documents.extend(PyPDFLoader(file_path).load())
        elif ext in csv_extensions:
          documents.extend(CSVLoader(file_path).load())
        elif ext in doc_extensions:
          try:
            documents.extend(UnstructuredWordDocumentLoader(file_path).load())
          except:
            try:
              import docx
              doc = docx.Document(file_path)
              text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
              documents.append(Document(page_content=text, metadata={"source": file_path, "type": "docx"}))
            except:
              print(f"[WARNING] Skipped {file_path} - could not load")
        elif ext in image_extensions:
          try:
            documents.extend(extract_text_from_image(file_path))
          except:
            print(f"[WARNING] Skipped {file_path} - could not extract text from image")
  
  print(f"[DEBUG] Loaded {len(documents)} documents from {source_dir}")
  
  if not documents:
    raise ValueError("No documents found in the specified sources or failed to extract text")
  
  return documents


def create_vector_store(docs: List[Document], embeddings, chunk_size: int = 10000, chunk_overlap: int = 200):
  """
  Create vector store from documents
  """
  text_splitter = RecursiveCharacterTextSplitter(
      chunk_size=chunk_size,
      chunk_overlap=chunk_overlap
  )
  splits = text_splitter.split_documents(docs)
  
  if not splits:
    raise ValueError("Documents did not contain any readable text chunks")
  
  # Use FAISS for retrieval
  return FAISS.from_documents(splits, embeddings).as_retriever(search_kwargs={"k": 5})


PROMPT_TEMPLATE = """
  Role & Scope
        You are HuduAssist, a real-time, authoritative, Kenyan Government information assistant.
        Your sole function is to assist users with verified, publicly available, real-time, and historical information about:
        - The Government of Kenya (ministries, departments, agencies, state corporations, and county governments).
        - Official Kenyan Government services, portals, and regulations.
        - Public announcements, policies, directives, and service procedures within Kenyan jurisdiction.
        You must not provide information unrelated to the Kenyan Government. If the query is outside your scope, politely decline and redirect the user.

        **Core Real-Time Functionality**
        Scrape & Fetch Live Data from the official list of government websites below at the moment of each request.
        Always prioritize the most relevant and official source for the query.
        Integrate breaking news, trending updates, and historical records for complete context.
        Continuously refresh data for high-traffic ministries and service portals to maintain accuracy.
        Maintain search awareness across multiple ministries simultaneously to provide a consolidated and authoritative answer.

        **High-Priority Sources - Main Gateway (Top Priority)**
https://www.hudumakenya.go.ke/

        **Ministries & Departments**
        - https://gok.kenya.go.ke/ministries
        - https://www.mod.go.ke/
        - https://www.ict.go.ke/
        - https://www.treasury.go.ke/
        - https://www.mfa.go.ke/
        - https://www.transport.go.ke/
        - https://www.lands.go.ke/
        - https://www.health.go.ke/
        - https://www.education.go.ke/
        - https://kilimo.go.ke/
        - https://www.trade.go.ke/
        - https://sportsheritage.go.ke/
        - https://www.environment.go.ke/
        - https://www.tourism.go.ke/
        - https://www.water.go.ke/
        - https://www.energy.go.ke/
        - https://www.labour.go.ke/
        - https://www.statelaw.go.ke/
        - https://www.president.go.ke/

        County Government
        - https://nairobi.go.ke/ and other county government sites.

        **Government Services Portals**
        - https://www.kra.go.ke/
        - https://www.kplc.co.ke/
        - https://accounts.ecitizen.go.ke/en
        - https://ardhisasa.lands.go.ke/home
        - https://teachersonline.tsc.go.ke/
        - https://sha.go.ke/

        **Response Guidelines**
        Always perform a live search/scrape of the relevant official site(s) before responding.
        Provide short, direct, conversational answers with the latest confirmed details.
Include a clickable link to the original source if requested or if procedural/legal/policy-related.
Do not guess — only provide verifiable facts.

  {context}

  Question: {question}
Answer:
"""


def get_qa_chain(source_dir):
  """
  Create QA chain with proper error handling
  """
  try:
    docs = load_documents(source_dir)
    llm, embeddings = load_model()
    retriever = create_vector_store(docs, embeddings)

    prompt = PromptTemplate(
      template=PROMPT_TEMPLATE,
      input_variables=["context", "question"]
    )

    def run_chain(inputs):
      query = inputs.get("query") or inputs.get("question")
      if not query:
        raise ValueError("A 'query' input is required")

      source_documents = retriever.invoke(query)
      context = "\n\n".join(doc.page_content for doc in source_documents)
      prompt_text = prompt.format(context=context, question=query)
      raw_response = llm.invoke(prompt_text)

      if hasattr(raw_response, "content"):
        answer = raw_response.content
      else:
        answer = raw_response

      if isinstance(answer, list):
        answer = " ".join(str(chunk) for chunk in answer)

      return {"result": answer, "source_documents": source_documents}

    return run_chain

  except Exception as e:
    print(f"[ERROR] Initializing QA system: {e}")
    return None


def query_system(query: str, qa_chain):
  """
  Query the QA system
  """
  if not qa_chain:
    return "System not initialized properly"

  try:
    result = qa_chain({"query": query})
    if not result["result"] or "don't know" in result["result"].lower():
      return "The answer could not be found in the provided documents"
    return f"HuduAssist 🇰🇪: {result['result']}"  
  except Exception as e:
    return f"Error processing query: {e}"


# Example usage:
# content_dir = "example.pdf"
# qa_chain = get_qa_chain(content_dir)
# print(query_system("What are the latest KRA services?", qa_chain))
